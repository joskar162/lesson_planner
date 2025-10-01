from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .forms import RegisterForm
from .models import LessonPlan
from django.http import HttpResponse, Http404
import io
from django.utils.crypto import get_random_string
from django.core.mail import send_mail
from django.utils import timezone
from .models import PasswordResetCode
from django.contrib.auth.models import User


def infer_student_requirements(topic: str) -> str:
    """Return a short suggested list of student requirements based on the topic."""
    if not topic:
        return ""
    t = topic.lower()
    # Simple keyword-based inference. Expand as needed.
    if "geometry" in t or "shapes" in t or "angle" in t or "triang" in t:
        return "Ruler, protractor, compass, calculator"
    if "algebra" in t or "equation" in t or "express" in t:
        return "Scientific calculator, algebraic notation familiarity"
    if "fractions" in t or "decimal" in t:
        return "Basic multiplication/division skills, fraction strips (optional)"
    if "probability" in t or "statistics" in t:
        return "Dice/coins or sample data, calculator"
    if "python" in t or "program" in t or "coding" in t:
        return "Laptop with Python installed or access to an online REPL, basic typing skills"
    if "binar" in t or "binary" in t:
        return "Paper and pencil for conversions, basic understanding of place value"
    if "chemistry" in t or "experiment" in t or "lab" in t:
        return "Lab coat, safety goggles, basic lab safety knowledge"
    if "history" in t or "geography" in t:
        return "Map or timeline materials, prior knowledge of relevant events"
    # default
    return "Basic classroom materials (notebook, pen), any prior prerequisite knowledge stated in unit overview"


def _generate_code():
    return get_random_string(20)


def forgot_password(request):
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            messages.error(request, 'No account found with that email address.')
            return render(request, 'forgot_password.html')

        code = _generate_code()
        PasswordResetCode.objects.create(user=user, code=code)

        # Send email (console backend in dev will print to console)
        subject = 'Your lesson planner password reset code'
        message = f'Use this code to reset your password: {code}\nOr follow the link: {request.build_absolute_uri("/reset-password/")}'
        send_mail(subject, message, None, [email], fail_silently=True)

        messages.success(request, 'A reset code has been sent to your email (check console in development).')
        return redirect('login')
    return render(request, 'forgot_password.html')


def reset_password(request):
    if request.method == 'POST':
        email = request.POST.get('email', '').strip()
        code = request.POST.get('code', '').strip()
        new_password = request.POST.get('new_password', '').strip()

        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            messages.error(request, 'Invalid email or code.')
            return render(request, 'reset_password.html')

        try:
            prc = PasswordResetCode.objects.filter(user=user, code=code).latest('created')
        except PasswordResetCode.DoesNotExist:
            messages.error(request, 'Invalid code.')
            return render(request, 'reset_password.html')

        # Optionally expire codes after some time
        if timezone.now() - prc.created > timezone.timedelta(hours=1):
            messages.error(request, 'Code expired.')
            return render(request, 'reset_password.html')

        if not new_password:
            messages.error(request, 'Enter a new password.')
            return render(request, 'reset_password.html')

        user.set_password(new_password)
        user.save()
        messages.success(request, 'Password reset successful. You can now log in.')
        return redirect('login')

    return render(request, 'reset_password.html')


def _make_pdf_bytes(text: str) -> bytes:
    """Return PDF bytes for the provided text using reportlab.

    Raises ImportError if reportlab is not installed.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        raise

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    textobject = p.beginText(40, 750)
    textobject.setFont("Helvetica", 11)
    for line in text.splitlines():
        if len(line) <= 90:
            textobject.textLine(line)
        else:
            for i in range(0, len(line), 90):
                textobject.textLine(line[i:i+90])
    p.drawText(textobject)
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.getvalue()


def _make_docx_bytes(text: str, lp=None) -> bytes:
    """Return DOCX bytes for the provided text. Uses python-docx.

    Raises ImportError if python-docx is not installed.
    """
    try:
        from docx import Document
    except ImportError:
        raise

    doc = Document()
    if lp is not None:
        doc.add_heading(f'Lesson Plan: {lp.topic}', level=1)
        doc.add_paragraph(f'Subject: {lp.subject}')
        doc.add_paragraph(f'Grade: {lp.grade}')
        doc.add_paragraph(f'Duration: {lp.duration} minutes')
        doc.add_paragraph('')
    for line in text.splitlines():
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _infer_student_requirements(topic: str) -> str:
    """Return a short string listing suggested student requirements based on topic keywords.

    This is a simple heuristic: match keywords in the topic to common supplies/knowledge.
    Extend this mapping as needed.
    """
    if not topic:
        return ''
    t = topic.lower()
    suggestions = []
    mapping = {
        'geometry': ['ruler', 'protractor', 'compass', 'mathematical set'],
        'algebra': ['calculator', 'pen', 'notebook'],
        'calculus': ['graphing calculator', 'notebook'],
        'statistics': ['calculator', 'spreadsheet (optional)'],
        'chemistry': ['safety goggles', 'lab coat', 'gloves'],
        'physics': ['calculator', 'meter stick', 'stopwatch'],
        'biology': ['lab coat', 'gloves', 'microscope (if available)'],
        'english': ['textbook', 'notebook', 'pen'],
        'history': ['textbook', 'timeline handout'],
        'art': ['pencils', 'eraser', 'colours', 'paper'],
        'music': ['instrument (if applicable)'],
        'programming': ['laptop', 'internet access'],
        'computer': ['laptop', 'internet access'],
        'math': ['calculator', 'notebook', 'pencil'],
    }
    for key, items in mapping.items():
        if key in t:
            suggestions.extend(items)

    # If none matched, provide a short generic suggestion
    if not suggestions:
        return 'Basic stationery (pen/pencil, notebook)'
    # deduplicate while preserving order
    seen = set()
    dedup = []
    for s in suggestions:
        if s not in seen:
            dedup.append(s)
            seen.add(s)
    return ', '.join(dedup)

@login_required
def index(request):
    lesson_plan_text = None
    lesson_plan_id = None
    # ensure these locals exist for GET requests
    topic = ''
    student_requirements = ''
    teacher_actions = ''
    duration_int = None
    if request.method == "POST":
        subject = request.POST.get("subject", "").strip()
        grade = request.POST.get("grade", "").strip()
        topic = request.POST.get("topic", "").strip()
        duration = request.POST.get("duration", "").strip()
        teacher_actions = request.POST.get("teacher_actions", "").strip()
        student_requirements = request.POST.get("student_requirements", "").strip()

        if not (subject and grade and topic and duration):
            messages.error(request, "All fields are required.")
        else:
            try:
                duration_int = int(duration)
            except ValueError:
                duration_int = None
                messages.error(request, "Duration must be a number.")

            if duration_int is not None:
                # Build plan and include teacher actions section
                lesson_plan_text = (
                    f"Subject: {subject}\n"
                    f"Grade: {grade}\n"
                    f"Topic: {topic}\n"
                    f"Duration: {duration_int} minutes\n\n"
                    "Objective:\n"
                    f"- Students will learn the basics of {topic}.\n\n"
                    "Materials:\n"
                    "- Whiteboard, markers, worksheets.\n"
                    f"- Student requirements: {student_requirements or 'None specified.'}\n\n"
                    "Activities (with teacher actions):\n"
                    "1. Introduction (10 min): Hook + objectives.\n"
                    f"   Teacher actions: {teacher_actions or 'Introduce topic, set objectives.'}\n\n"
                    "2. Teaching & Modelling:\n"
                    f"   Teacher actions: {teacher_actions or 'Explain and model examples.'}\n\n"
                    "3. Guided Practice:\n"
                    f"   Teacher actions: {teacher_actions or 'Guide students through examples.'}\n\n"
                    "4. Independent Practice:\n"
                    f"   Teacher actions: {teacher_actions or 'Monitor and support.'}\n\n"
                    "5. Assessment & Plenary:\n"
                    f"   Teacher actions: {teacher_actions or 'Give quick quiz and recap.'}\n\n"
                    "Homework:\n"
                    f"- Practice problems on {topic}."
                )

                lp = LessonPlan.objects.create(
                    user=request.user,
                    subject=subject,
                    grade=grade,
                    topic=topic,
                    duration=duration_int,
                    content=lesson_plan_text,
                    teacher_actions=teacher_actions,
                    student_requirements=student_requirements,
                )
                lesson_plan_id = lp.id
                messages.success(request, "Lesson plan generated and saved.")

                # If the user clicked Generate & Download PDF/DOCX, return the generated file immediately
                if 'download_pdf' in request.POST:
                    try:
                        pdf_bytes = _make_pdf_bytes(lesson_plan_text)
                        response = HttpResponse(pdf_bytes, content_type='application/pdf')
                        response['Content-Disposition'] = f'attachment; filename="lessonplan_{lp.id}.pdf"'
                        response['Content-Length'] = str(len(pdf_bytes))
                        return response
                    except ImportError:
                        # Fallback: return plain text file with only the generated part
                        txt = lesson_plan_text or ''
                        txt_bytes = txt.encode('utf-8')
                        response = HttpResponse(txt_bytes, content_type='text/plain; charset=utf-8')
                        response['Content-Disposition'] = f'attachment; filename="lessonplan_{lp.id}.txt"'
                        response['Content-Length'] = str(len(txt_bytes))
                        return response

                if 'download_docx' in request.POST:
                    try:
                        docx_bytes = _make_docx_bytes(lesson_plan_text, lp)
                        response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename="lessonplan_{lp.id}.docx"'
                        response['Content-Length'] = str(len(docx_bytes))
                        return response
                    except ImportError:
                        # Fallback to plain text attachment
                        txt = lesson_plan_text or ''
                        txt_bytes = txt.encode('utf-8')
                        response = HttpResponse(txt_bytes, content_type='text/plain; charset=utf-8')
                        response['Content-Disposition'] = f'attachment; filename="lessonplan_{lp.id}.txt"'
                        response['Content-Length'] = str(len(txt_bytes))
                        return response

    recent_plans = LessonPlan.objects.filter(user=request.user).order_by("-created")[:10]
    # If student_requirements wasn't provided, infer from topic for display
    display_student_requirements = student_requirements if (student_requirements is not None and student_requirements != '') else (infer_student_requirements(topic) if topic else '')
    return render(request, "index.html", {"lesson_plan": lesson_plan_text, "recent_plans": recent_plans, "lesson_plan_id": lesson_plan_id, "student_requirements": display_student_requirements})

def welcome(request):
    return render(request, 'welcome.html')

def register(request):
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            form.save()                      # saves user with hashed password
            messages.success(request, "Account created. Please log in.")
            return redirect('login')         # go to login page after register
        else:
            messages.error(request, "Registration failed. Fix the errors below.")
            print("Registration errors:", form.errors)  # check runserver console
    else:
        form = RegisterForm()
    return render(request, 'register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
            return redirect('home')
        else:
            return render(request, 'login.html', {'error': 'Invalid credentials'})
    return render(request, 'login.html')

def logout_view(request):
    logout(request)
    return redirect('welcome')


@login_required
def lesson_pdf(request, pk):
    """Generate a PDF for the requested LessonPlan and return it as attachment.

    If reportlab is not installed, redirect back with an instructive message.
    """
    try:
        lp = LessonPlan.objects.get(pk=pk, user=request.user)
    except LessonPlan.DoesNotExist:
        raise Http404("Lesson plan not found")

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        # Fallback: return plain text attachment with the lesson plan content
        txt = lp.content or ''
        txt_bytes = txt.encode('utf-8')
        response = HttpResponse(txt_bytes, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="lessonplan_{pk}.txt"'
        response['Content-Length'] = str(len(txt_bytes))
        return response

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    textobject = p.beginText(40, 750)
    textobject.setFont("Helvetica", 11)
    for line in lp.content.splitlines():
        # Basic text wrapping for very long lines
        if len(line) <= 90:
            textobject.textLine(line)
        else:
            # naive wrap: split into chunks of ~90 chars
            for i in range(0, len(line), 90):
                textobject.textLine(line[i:i+90])
    p.drawText(textobject)
    p.showPage()
    p.save()
    buffer.seek(0)
    pdf_bytes = buffer.getvalue()

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="lessonplan_{pk}.pdf"'
    response['Content-Length'] = str(len(pdf_bytes))
    return response



@login_required
def lesson_docx(request, pk):
    """Generate a .docx (Word) document for the requested LessonPlan and return it as attachment.

    If python-docx is not installed, redirect back with an instructive message.
    """
    try:
        lp = LessonPlan.objects.get(pk=pk, user=request.user)
    except LessonPlan.DoesNotExist:
        raise Http404("Lesson plan not found")

    try:
        from docx import Document
    except ImportError:
        # Fallback: return plain text attachment with the lesson plan content
        txt = lp.content or ''
        txt_bytes = txt.encode('utf-8')
        response = HttpResponse(txt_bytes, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="lessonplan_{pk}.txt"'
        response['Content-Length'] = str(len(txt_bytes))
        return response

    doc = Document()
    doc.add_heading(f'Lesson Plan: {lp.topic}', level=1)
    # Add metadata lines as paragraphs for better Word formatting
    doc.add_paragraph(f'Subject: {lp.subject}')
    doc.add_paragraph(f'Grade: {lp.grade}')
    doc.add_paragraph(f'Duration: {lp.duration} minutes')
    doc.add_paragraph('')
    for line in lp.content.splitlines():
        # skip empty lines to avoid excessive spacing
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.getvalue()

    response = HttpResponse(docx_bytes, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="lessonplan_{pk}.docx"'
    response['Content-Length'] = str(len(docx_bytes))
    return response

